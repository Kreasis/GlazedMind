"""Rebuild the static Knowledge Base directly in PostgreSQL/pgvector."""

import re
import sys
from pathlib import Path

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from pypdf import PdfReader

BASE_DIR = Path(__file__).resolve().parents[1]
if str(BASE_DIR) not in sys.path:
    sys.path.insert(0, str(BASE_DIR))
KNOWLEDGE_DIR = BASE_DIR / "knowledge-base"
IMAGE_DIR = BASE_DIR / "data" / "knowledge_images"

def slugify(value: str) -> str:
    return re.sub(r"[^a-z0-9]+", "-", value.lower()).strip("-")

def _extension(part: object) -> str:
    suffix = Path(str(getattr(part, "partname", ""))).suffix.lower()
    return suffix if suffix else ".png"

def _paragraph_images(paragraph: Paragraph, document: Document, target: Path, slug: str, start: int) -> tuple[list[str], int]:
    images: list[str] = []
    index = start
    for relationship_id in paragraph._p.xpath(".//a:blip/@r:embed"):
        part = document.part.related_parts.get(relationship_id)
        if part is None or not hasattr(part, "blob"):
            continue
        index += 1
        filename = f"{index}{_extension(part)}"
        (target / filename).write_bytes(part.blob)
        images.append(f"{slug}/{filename}")
    return images, index

def extract_docx(path: Path, image_root: Path | None = None) -> tuple[str, list[dict[str, object]], list[str]]:
    """Extract paragraphs and screenshots in their actual Word document order."""
    document = Document(path)
    tagged = (
        any(re.match(r"^SECTION\s*:", paragraph.text.strip(), re.I) for paragraph in document.paragraphs)
        and any(re.match(r"^STEP\s+\d+[A-Z]?\s*$", paragraph.text.strip(), re.I) for paragraph in document.paragraphs)
    )
    slug = slugify(path.stem)
    target = (image_root or IMAGE_DIR) / slugify(path.stem)
    target.mkdir(parents=True, exist_ok=True)
    records: list[dict[str, object]] = []
    images: list[str] = []
    image_index = 0
    all_text: list[str] = []
    current_section = "Procedure"
    pending_step: dict[str, object] | None = None

    def finish_step() -> None:
        nonlocal pending_step
        if pending_step and str(pending_step.get("text", "")).strip():
            records.append(pending_step)
        pending_step = None

    def consume(paragraph: Paragraph) -> None:
        nonlocal image_index, current_section, pending_step
        text = paragraph.text.strip()
        paragraph_images, image_index = _paragraph_images(paragraph, document, target, slug, image_index)
        images.extend(paragraph_images)
        if text:
            all_text.append(text)
        if tagged:
            section_match = re.match(r"^SECTION\s*:\s*(.+)$", text, re.I)
            step_match = re.match(r"^STEP\s+(\d+[A-Z]?)\s*$", text, re.I)
            notice_match = re.match(r"^(NOTE|WARNING)\s*(.*)$", text, re.I | re.S)
            if section_match:
                finish_step()
                current_section = section_match.group(1).strip()
                return
            if step_match:
                finish_step()
                pending_step = {"text": "", "images": [], "section": current_section, "kind": "step"}
                return
            if notice_match:
                finish_step()
                notice_text = notice_match.group(2).strip()
                if notice_text:
                    records.append({"text": notice_text, "images": paragraph_images, "section": current_section, "kind": notice_match.group(1).lower()})
                return
            if pending_step is not None:
                if text:
                    existing = str(pending_step["text"])
                    pending_step["text"] = f"{existing} {text}".strip()
                pending_step["images"].extend(paragraph_images)
            return
        if text:
            records.append({"text": text, "images": paragraph_images})
        elif paragraph_images and records:
            # Screenshots placed immediately below an instruction belong to it.
            records[-1]["images"].extend(paragraph_images)

    for block in document.iter_inner_content():
        if isinstance(block, Paragraph):
            consume(block)
        elif isinstance(block, Table):
            for row in block.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        consume(paragraph)
    finish_step()
    content = "\n".join(all_text if tagged else (str(record["text"]) for record in records))
    return content, records, images

def read_document(path: Path) -> tuple[str, list[str]]:
    if path.suffix.lower() == ".pdf":
        text = "\n".join(page.extract_text() or "" for page in PdfReader(path).pages)
        return text, []
    return "", []

def main() -> None:
    from app.services.vector_store import replace_documents

    documents = []
    for path in sorted(KNOWLEDGE_DIR.iterdir()):
        if path.suffix.lower() not in {".docx", ".pdf"}:
            continue
        if path.suffix.lower() == ".docx":
            raw_text, steps, images = extract_docx(path)
        else:
            raw_text, steps = read_document(path)
            images = []
        text = " ".join(raw_text.split())
        documents.append({"title": path.stem, "filename": path.name, "content": text, "steps": steps, "images": images})
    count = replace_documents(documents)
    print(f"Indexed {count} knowledge documents in PostgreSQL/pgvector")

if __name__ == "__main__":
    main()
